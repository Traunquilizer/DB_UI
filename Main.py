# -*- coding: utf-8 -*-

import re
import sys
import os
import pymongo
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from bson.objectid import ObjectId
import json

from name import *
from yu import *
from dialogwins import *
from lib import *

from PyQt5 import QtCore, QtGui, QtWidgets


def server():
    # путь прописывается индивидуально, либо
    # при значении пути по умолчанию аргумент опускается
    os.system('mongodb' + ' --dbpath /home/traun/mongodb')

def client():
    client = pymongo.MongoClient()
    db = client.Architecht

config = {
    'Принято': {
        'db': db.Applications,
        'id_f': 10,
        'header': ['Дата', 'ФИО', 'Телефон', 'Изделие', 'Примечания', 
        'Принял', 'Место', 'Гарантийность', 'Производитель','Акт приёма', 
        'id'],
        'filling_func': lambda i: [str(i['Дата'])[8:] + '/' +
                                    str(i['Дата'])[5:7] + '/' +
                                    str(i['Дата'])[0:4], 
                                    i['ФИО'], 
                                    i['Телефон'],
                                    i['Изделие'],
                                    i['Примечания'],
                                    i['Принял'],
                                    i['Место'],
                                    i['Гарантийность'],
                                    i['Производитель'],
                                    i['Акт приёма'],
                                    str(i['_id'])],
        'delegate_col': [4, 6],
        'editing_fs': { 4: {'editor': 'plain_text',
                            'field': 'Примечания'},
                        6: {'editor': 'combo',
                            'field': 'Место',
                            'c_source': [db.Places, 'Место']}},
        'default_argument':  lambda i: {'$text' : { '$search': i }},
        'Status': 'Принято',
        'sorting_criteria': ['Дата', 'ФИО', 'Телефон', 'Изделие', 'Примечания', 
        'Принял', 'Место', 'Гарантийность', 'Производитель', 'Акт приёма']
                },
    'В работе': {
        'db': db.Applications,
        'id_f': 7,
        'header': ['ФИО', 'Изделие', 'Примечания', 'Место', 'Кто передал в ремонт', 
        'Работнику', 'Перечень работ', 'id'],
        'filling_func': lambda i: [i['ФИО'], 
                                    i['Изделие'],
                                    i['Примечания'],
                                    i['Место'], 
                                    i['Передал в ремонт'],
                                    i['Передано работнику'],
                                    i['Перечень работ'],
                                    str(i['_id'])],
        'default_argument':  lambda i: {'$text' : { '$search': i }},
        'Status': 'В работе',
        'delegate_col': [2, 3, 4, 5, 6],
        'editing_fs': { 2: {'editor': 'plain_text',
                            'field': 'Примечания'},
                        3: {'editor': 'combo',
                            'c_source': [db.Places, 'Место'],
                            'field': 'Место'},
                        4: {'editor': 'combo',
                            'c_source': [db.Receivers, 'ФИО'],
                            'field': 'Передал в ремонт'},
                        5: {'editor': 'combo',
                            'c_source': [db.Workers, 'ФИО'],
                            'field': 'Передано работнику'},
                        6: {'editor': 'c_line',
                            'c_source': [db.Jobs, 'Наименование'],
                            'cu_income': 'jobs',
                            'field': 'Перечень работ'}},
        'sorting_criteria': ['Дата', 'Изделие', 'Примечания', 'Место',
        'Передал в ремонт', 'Передано работнику', '_id']
                },
    'Завершено': {
        'db': db.Applications,
        'id_f': 8,
        'header': ['ФИО', 'Изделие', 'Место', 'Перечень работ', 'Дата завершения', 
        'Цена', 'Зарплата', 'Расчёт', 'id'],
        'filling_func': lambda i: [i['ФИО'],
                                    i['Изделие'],
                                    i['Место'], 
                                    i['Перечень работ'],
                                    str(i['Дата завершения'])[8:] + '/' +
                                    str(i['Дата завершения'])[5:7] + '/' +
                                    str(i['Дата завершения'])[0:4],
                                    str(i['Цена']),
                                    str(i['Зарплата']),
                                    str(i['Расчёт']),
                                    str(i['_id'])],
        'delegate_col': [2, 4, 7],
        'editing_fs': { 2: {'editor': 'combo',
                            'c_source': [db.Places, 'Место'],
                            'field': 'Место'},
                        4: {'editor': 'date',
                            'cu_income': 'date',
                            'field': 'Дата завершения'},
                        7: {'editor': 'line',
                            'field': 'Расчёт',
                            'cu_income': 'price'}},
        'default_argument':  lambda i: {'$text' : { '$search': i }},
        'Status': 'Завершено',
        'sorting_criteria': ['Дата', 'ФИО', 'Изделие', 'Место', 'Перечень работ', 
        'Дата завершения', 'Цена', 'Зарплата', 'Расчёт', '_id']
                },
    'Архив': {
        'db': db.Applications,
        'id_f': 8,
        'header': ['ФИО', 'Изделие', 'Дата приёма', 'Примечания', 'Получатель', 
        'Цена', 'Расчёт', 'Дата отгрузки', 'id'],
        'filling_func': lambda i: [i['ФИО'],
                                    i['Изделие'],
                                    i['Дата'],
                                    i['Примечания'],
                                    i['Получатель'],
                                    str(i['Цена']),
                                    str(i['Расчёт']),
                                    str(i['Дата отгрузки'])[8:] + '/' +
                                    str(i['Дата отгрузки'])[5:7] + '/' +
                                    str(i['Дата отгрузки'])[0:4],
                                    str(i['_id'])],
        'delegate_col': [3, 4, 6, 7],
        'editing_fs': { 3: {'editor': 'plain_text',
                            'field': 'Примечания'},
                        4: {'editor': 'line',
                            'field': 'Получатель'},
                        6: {'editor': 'line',
                            'field': 'Расчёт',
                            'cu_income': 'price'},
                        7: {'editor': 'date',
                            'cu_income': 'date',
                            'field': 'Дата отгрузки'}},
        'default_argument':  lambda i: {'$text' : { '$search': i }},
        'Status': 'Архив',
        'sorting_criteria': ['Дата', 'ФИО', 'Изделие', 'Архивные примечания', 
        'Дата отгрузки', '_id']
                },
    'Клиенты': {
        'db': db.Applicants,
        'id_f': 2,
        'header': ['Клиент', 'Телефон', 'id'],
        'filling_func': lambda i: [i['Клиент'], '+7 ' + i['Телефон'], str(i['_id'])],
        'delegate_col': [0, 1],
        'editing_fs': { 0: {'editor': 'line',
                            'field': 'Клиент'},
                        1: {'editor': 'line',
                            'field': 'Телефон',
                            'cu_income': 'phone'}},
        'db_name': 'клиентов',
        'fields': { 'Клиент': { 'disp_field': 'Клиент',
                                'editor': 'line'},
                    'Телефон': {'disp_field': 'Телефон',
                                'editor': 'line',
                                'cu_income': 'phone'}},
        'fields_order': ['Клиент', 'Телефон'],
        'default_argument':  lambda i: {'$text' : { '$search': i }},
        'Status': 'active',
        'sorting_criteria': ['Клиент', 'Телефон', '_id']
                },
    'Работники': {
        'db': db.Workers,
        'id_f': 1,
        'header': ['ФИО', 'id'],
        'filling_func': lambda i: [i['ФИО'], str(i['_id'])],
        'delegate_col': [0],
        'editing_fs': { 0: {'editor': 'line',
                            'field': 'ФИО'}},
        'db_name': 'работников',
        'fields': {'ФИО': { 'disp_field': 'ФИО',
                            'editor': 'line'}},
        'fields_order': ['ФИО'],
        'default_argument':  lambda i: {'$text' : { '$search': i }},
        'Status': 'active',
        'sorting_criteria': ['ФИО', '_id']
                },
    'Приёмщики склада': {
        'db': db.Receivers,
        'id_f': 1,
        'header': ['ФИО', 'id'],
        'filling_func': lambda i: [i['ФИО'], str(i['_id'])],
        'delegate_col': [0],
        'editing_fs': { 0: {'editor': 'line',
                            'field': 'ФИО'}},
        'db_name': 'приёмщиков',
        'fields': {'ФИО': { 'disp_field': 'ФИО',
                            'editor': 'line',
                            'cu_income': 'names'}},
        'fields_order': ['ФИО'],
        'default_argument':  lambda i: {'$text' : { '$search': i }},
        'Status': 'active',
        'sorting_criteria': ['ФИО', '_id']
                        },
    'Операции': {
        'db': db.Jobs,
        'id_f': 3,
        'header': ['Наименование', 'Стоимость', 'Оплата работнику', 'id'],
        'filling_func': lambda i: [i['Наименование'], i['Стоимость'],  
        i['Оплата работнику'], str(i['_id'])],
        'delegate_col': [0, 1, 2],
        'editing_fs': { 0: {'editor': 'line',
                            'field': 'Наименование'},
                        1: {'editor': 'line',
                            'field': 'Стоимость',
                            'cu_income': 'price'},
                        2: {'editor': 'line',
                            'field': 'Оплата работнику',
                            'cu_income': 'price'}},
        'db_name': 'операций',
        'fields': { 'Наименование':{'disp_field': 'Наименование',
                                    'editor': 'line',},
                    'Стоимость':   {'disp_field': 'Стоимость',
                                    'cu_income': 'price',
                                    'editor': 'line'},
                    'Оплата работнику':{'disp_field':'Оплата работнику',
                                    'cu_income': 'price',
                                    'editor': 'line'}},
        'fields_order':['Наименование','Стоимость','Оплата работнику'],
        'default_argument':  lambda i: {'$text' : { '$search': i }},
        'Status': 'active',
        'sorting_criteria': ['Наименование', 'Стоимость', '_id']
                },
    'Изделия': {
        'db': db.Prods,
        'id_f': 1,
        'header': ['Изделие', 'id'],
        'filling_func': lambda i: [i['Изделие'], str(i['_id'])],
        'delegate_col': [0],
        'editing_fs': { 0: {'editor': 'line',
                            'field': 'Изделие'}},
        'db_name': 'изделий',
        'fields': {'Изделие': { 'disp_field': 'Наименование',
                                'editor': 'line'}},
        'fields_order': ['Изделие'],
        'default_argument':  lambda i: {'$text' : { '$search': i }},
        'Status': 'active',
        'sorting_criteria': ['Изделие', '_id']
                },
    'Место': {
        'db': db.Places,
        'id_f': 1,
        'header': ['Место', 'id'],
        'filling_func': lambda i: [i['Место'], str(i['_id'])],
        'delegate_col': [0],
        'editing_fs': {0 : {'editor': 'line',
                            'field': 'Место'}},
        'db_name': 'мест',
        'fields': {'Место': {'disp_field': 'Место',
                             'editor': 'line'}},
        'fields_order': ['Место'],
        'default_argument':  lambda i: {'$text' : { '$search': i }},
        'Status': 'active',
        'sorting_criteria': ['Место', '_id']
                },
    'Акты приёма в ремонт': {
        'db': db.Checks,
        'id_f': 8,
        'header': ['Изделия', 'Гарантийность', 'Принято от', 'Контакт', 
        'Примечания', 'Принял', 'Дата', 'Num', 'id'],
        'filling_func': lambda i: [i['Изделия'], i['Гарантийность'], 
        i['Принято от'], i['Контакт'], i['Примечания'], i['Принял'], 
        str(i['Дата'])[8:] + '/' + str(i['Дата'])[5:7] + '/' + str(i['Дата'])[0:4], 
        i['Num'], str(i['_id'])],
        'delegate_col': [],
        'fields_order': [],
        'default_argument':  lambda i: {'$text' : { '$search': i }},
        'Status': 'active',
        'sorting_criteria': ['Изделия', 'Гарантийность', 'Принято от', 
        'Контакт', 'Примечания', 'Принял', 'Дата', 'Num', '_id']
                }
         }

conf_crf = {'line':lambda parent, source_list = None: QtWidgets.QLineEdit(parent),
            'plain_text': lambda parent, source_list = None: QtWidgets.QPlainTextEdit(parent),
            'combo': lambda parent, source_list: combo_editor(source_list, parent),
            'date': lambda parent, source_list = None: date_editor(parent),
            'c_line': lambda parent, source_list: create_c_line(parent, source_list)}
conf_set = {'line':lambda editor, index: editor.setText(index.data()),
            'plain_text': lambda editor, index: editor.setPlainText(index.data()),
            'combo': lambda editor, index: combo_editor_set(editor),
            'date': lambda editor, index: editor.setCurrentSection(QtWidgets.QDateTimeEdit.MonthSection)}
conf_stm = {'line': lambda editor: editor.text(),
            'plain_text': lambda editor: editor.text(),
            'combo': lambda editor: editor.currentText(),
            'date': lambda editor: editor.date()}
cu_conf =  {'price': lambda data, ident = None: new_data_prices(data),
            'date': lambda data, ident = None: new_data_date(data),
            'jobs': lambda data, ident: new_data_jobs(data, ident),
            'phone': lambda data, ident = None: new_data_phones(data),
            'names': lambda data, ident = None: new_data_names(data)}


def combo_editor_set(editor):
    editor.blockSignals(True)
    editor.setCurrentIndex(0)
    editor.blockSignals(False)

def combo_editor(source_list, parent):
    combo = QtWidgets.QComboBox(parent)
    combo.addItems(get_source_data(source_list[0], source_list[1]))
    return combo

def set_completer(widget, source_list):
    completer = Completer(get_source_data(source_list[0], source_list[1]))
    widget.setCompleter(completer)
    return widget

def create_c_line(parent, source_list):
    line = QtWidgets.QLineEdit(parent)
    return set_completer(line, source_list)

def date_editor(parent):
    dateEdit = QtWidgets.QDateEdit(parent)
    dateEdit.setAutoFillBackground(False)
    dateEdit.setCurrentSection(QtWidgets.QDateTimeEdit.MonthSection)
    dateEdit.setDateTime(QtCore.QDateTime.currentDateTime())
    dateEdit.setCalendarPopup(True)
    return dateEdit

def get_source_data(db, field):
    func = db.find
    data = []
    for i in func():
        data.append(i[field])
    return data

def new_data_date(new_data):
    new_data = str(new_data.toPyDate())
    return new_data

def new_data_jobs(new_data, ident):
    price = 0
    salary = 0
    dictof = sep_count(new_data)
    for i in dictof:
        if not db.Jobs.find_one({'Наименование': i}):
            QtWidgets.QMessageBox.warning(sew, texts_conf["error"]["title"], texts_conf["error"]["text"]["jobs"])
            return ''
        else:
            price += int(float(db.Jobs.find_one({'Наименование': i})['Стоимость']))*dictof[i][0]
            salary += int(float(db.Jobs.find_one({'Наименование': i})['Оплата работнику']))*dictof[i][0]
    db.Applications.find_one_and_update({'_id': ObjectId(ident)},
            {'$set': {'Цена': price, 'Зарплата': salary}})
    return new_data

def new_data_phones(new_data):
    if not re.match(r'[0-9]', new_data):
        QtWidgets.QMessageBox.warning(sew, texts_conf["error"]["title"] , texts_conf["error"]["phones"])
        return ''
    elif not len(new_data) in [11, 12]:
        if not new_data[0:2]=='+7' or new_data[0] == '8' and len(new_data)==12:
            reply = QtWidgets.QMessageBox.warning(sew, texts_conf["warning"]["title"],
                texts_conf["warning"]["text"]["phones"].format(new_data),
                QtWidgets.QMessageBox.Yes, QtWidgets.QMessageBox.No)
            if reply == QtWidgets.QMessageBox.Yes:
                return new_data
            if reply == QtWidgets.QMessageBox.No:
                return ''
        else:
            return new_data
    else:
        return new_data

def new_data_prices(data):
    try:
        float(data)
        return str(data)
    except ValueError:
        QtWidgets.QMessageBox.warning(sew, texts_conf["error"]["title"] , 
            texts_conf["error"]["text"]["prices"])
        return None

def new_data_names(data):
    if not re.match(r'^\w+[ ]\w+?$', data):
        QtWidgets.QMessageBox.warning(sew, texts_conf["error"]["title"], 
            texts_conf["error"]["text"]["name"])
        return None
    else:
        return data

def str_date(date):
    dat = {'01':'января', '02':'февраля', '03':'марта', '04':'апреля', '05':'мая', '06':'июня', 
    '07':'июля', '08':'августа', '09':'сентября', '10':'октября', '11':'ноября', '12':'декабря'}
    str_date = '{} {} {}г'.format(date[8:],dat[date[5:7]],date[:4])
    return str_date


class MyWin(QtWidgets.QMainWindow):

    def __init__(self, parent=None):

        QtWidgets.QWidget.__init__(self, parent)
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.reset_completers()

        self.ui.button_proceed.clicked.connect(self.init_insert)
        self.ui.button_search.clicked.connect(self.search_func)

    def complete_phone(self):       #автодополнение телефона
        name = self.ui.line_name.text()
        try:
            phone = db.Applicants.find_one({'Клиент': name})['Телефон']
            self.ui.line_phone.setText(phone)
        except TypeError:
            pass

    def win_clear(self):                #очищение полей
        self.ui.line_product_name.setText('')
        self.ui.line_name.setText('')
        self.ui.textf_stuff.setPlainText('')
        self.ui.line_phone.setText('')
        self.ui.checkBox_guarantee.setChecked(False)
        self.reset_completers()

    def insert_func(self, check_num, product_name, date, name, phone, receiver, stuff, guarantee, producer):                                                            #ввод в базу
        db.Applications.insert_one({'ФИО': name,
                                    'Телефон': phone,
                                    'Принял': receiver,
                                    'Изделие': product_name,
                                    'Примечания': stuff,
                                    'Дата': date,
                                    'Производитель': producer,
                                    'Акт приёма': check_num,
                                    'Гарантийность': guarantee,
                                    'Статус': 'Принято',
                                    'Место': '',
                                    'Передал в ремонт': '',
                                    'Передано работнику': '',
                                    'Перечень работ': '',
                                    'Дата завершения': '',
                                    'Цена': '',
                                    'Дата отгрузки': '',
                                    'Зарплата': '',
                                    'Расчёт': '',
                                    'Получатель': ''})
        # sew.scaler()
        self.win_clear()
        self.reset_completers()
        QtWidgets.QMessageBox.information(self, 'Выполнено', 'Заявка успешно сохранена')

    def update_client(self, name, phone):                                                       #обновление записи клиента
        if db.Applicants.find_one({'Клиент': name}) and not db.Applicants.find_one({'Клиент': name, 
            'Телефон': phone}):
            temp_phone = db.Applicants.find_one({'Клиент': name})['Телефон']
            msg = texts_conf["message"]["text"]["upd_client"] .format(name, temp_phone)
            reply = QtWidgets.QMessageBox.question(self, texts_conf["message"]["title"], msg,
                                                   QtWidgets.QMessageBox.Yes,
                                                   QtWidgets.QMessageBox.No)
            if reply == QtWidgets.QMessageBox.Yes:
                db.Applicants.find_one_and_update({'Клиент': name}, {'$set': {'Телефон': phone}})
                self.reset_completers()
            elif reply == QtWidgets.QMessageBox.No:
                pass
        elif not db.Applicants.find_one({'Клиент': name}):
            db.Applicants.insert_one({'Клиент': name, 'Телефон': phone})
            self.reset_completers()

    def reset_completers(self):                                                     #перезапуск автодополнения
        self.list_name = get_source_data(db.Applicants, 'Клиент')
        self.list_phone = get_source_data(db.Applicants, 'Телефон')
        self.completer = QtWidgets.QCompleter(self.list_name)
        self.completer_p = QtWidgets.QCompleter(self.list_phone)
        self.ui.line_name.setCompleter(self.completer)
        self.ui.line_phone.setCompleter(self.completer_p)
        set_completer(self.ui.line_product_name, [db.Prods, 'Изделие'])

        self.ui.dateEdit.setDateTime(QtCore.QDateTime.currentDateTime())

        self.completer.activated.connect(self.complete_phone)
        self.ui.cBox_receiver.clear()
        _translate = QtCore.QCoreApplication.translate
        for i in list(enumerate(db.Receivers.find())):
            self.ui.cBox_receiver.addItem("")
        for i in list(enumerate(db.Receivers.find())):
            self.ui.cBox_receiver.setItemText(i[0], _translate("MainWindow", i[1]['ФИО']))

    def create_doc(self, heading, content_list, save_to, document=None):
        if document == None:
            document = Document()
        else:
            document.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH, builtin=True)
            document.styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH, builtin=True)
            document.styles.add_style('TableGrid', WD_STYLE_TYPE.TABLE, builtin=True)
        document.add_heading(heading, 0)

        def create_fieldline(document, field_data):
            field = field_data[0]
            data = str(field_data[1])
            first_line = document.add_paragraph('')
            first_line.add_run(field).bold = True
            first_line.add_run('  '+data)

        def create_table(document, header_contents):
            header = header_contents[0]
            contents = header_contents[1]
            table = document.add_table(rows=len(contents)+1, cols=len(header))
            table.style = 'TableGrid'
            hdr_cells = table.rows[0].cells
            cntr = 0
            for i in enumerate(header):
                hdr_cells[i[0]].text = i[1]
            for item in enumerate(contents):
                cntr+=1
                for x in enumerate(header):
                    table.rows[cntr].cells[x[0]].text = str(item[1][x[0]])

        conf = {'field': create_fieldline,
                'table': create_table}

        for i in content_list:
            conf[i['type']](document, i['content'])

        document.add_page_break()
        wd = os.getcwd()
        cwd = wd[:]
        for i in save_to.split('/'):
            try:
                os.chdir(cwd+'/'+i)
                cwd += '/'+i
            except FileNotFoundError:
                os.mkdir(cwd+'/'+i)
                os.chdir(cwd+'/'+i)
                cwd += '/'+i
        document.save(heading + '.docx')
        os.chdir(wd)
        return document

    def init_insert(self):                                                          #проверка корректности заполнения
        temp_var = self.ui.dateEdit.date()
        date = str(temp_var.toPyDate())
        product_name = r'{}'.format(self.ui.line_product_name.text())
        name = r'{}'.format(self.ui.line_name.text())
        receiver = r'{}'.format(str(self.ui.cBox_receiver.currentText()))
        stuff = r'{}'.format(self.ui.textf_stuff.toPlainText())
        if stuff == '':
            stuff = ' -- '
        phone = new_data_phones(r'{}'.format(self.ui.line_phone.text()))
        if self.ui.checkBox_guarantee.isChecked():
            guarantee = 'по гарантии'
        else:
            guarantee = 'не по гарантии'
        if not self.check_func(product_name, name, phone):
            return 1
        self.update_client(name, phone)
        prodc = sep_count(product_name)
        prodc = self.producer(prodc)
        prodstr = self.prod_str(prodc)
        check_num = db.Checks.count_documents({'year': date[:4]})+1
        if not self.form_check(date, name, phone, receiver, stuff, guarantee, prodstr, check_num):
            return 1
        db.Checks.insert_one({'Изделия': prodstr[1:], 'Гарантийность': guarantee, 'Принято от': name, 
            'Контакт': phone, 'Примечания': stuff, 'Принял': receiver, 'Дата': date, 'Статус': 'active', 
            'Num': check_num, 'year': date[:4]})
        self.reset_completers()
        gotlist = []
        for i in prodc:
            for each in range(prodc[i][0]):
                self.insert_func(check_num, i, date, name, phone, receiver, stuff, guarantee, prodc[i][1])
                gotlist.append(str(db.Applications.find_one({'ФИО': name,
                                    'Телефон': phone,
                                    'Принял': receiver,
                                    'Изделие': product_name,
                                    'Примечания': stuff,
                                    'Дата': date,
                                    'Производитель': prodc[i][1],
                                    'Акт приёма': check_num}, {'_id':1})))
        db.Checks.find_one_and_update({'Num': check_num }, {'$set': {'gotlist': gotlist}})
        return 0

    def producer(self, prod_count):
        for i in prod_count:
            if db.Prods.find_one({'Изделие': i}):
                prod_count[i].append('ДНТ')
            else:
                prod_count[i].append('Сторонний')
        return prod_count

    def prod_str(self, prod_count):
        prod_str = ''
        for i in prod_count:
            prod_str += '\n - {} {}шт. Производитель: {}'.format(i, prod_count[i][0], prod_count[i][1])
        return prod_str

    def form_check(self, date, name, phone, receiver, stuff, guarantee, prod_str, check_num):     
        msg = texts_conf["receive_act"]["text"].format(check_num, str_date(date), prod_str, guarantee, 
            name, phone, stuff, receiver, date)
        reply = QtWidgets.QMessageBox.question(self, texts_conf["receive_act"]["title"], msg,
            QtWidgets.QMessageBox.Yes,QtWidgets.QMessageBox.No)
        try:
            cost = db.Jobs.find_one({'Наименование': 'Хранение'}, {'Стоимость': 1})['Стоимость']
        except:
            cost = '50'
        if reply == QtWidgets.QMessageBox.Yes:
            self.create_doc( texts_conf["docstr"]["receive_act"]["title"] .format(check_num, str_date(date)) , 
                [{'content': ['Изделия', prod_str[1:]], 'type': 'field'},
                {'content': ['Гарантийность', guarantee], 'type': 'field'},
                {'content': ['Принято от',name], 'type': 'field'},
                {'content': ['Контакт', phone], 'type': 'field'},
                {'content': ['Примечания', stuff], 'type': 'field'},
                {'content': ['Принял',receiver], 'type': 'field'},
                {'content': ['Дата', str_date(date)], 'type': 'field'},
                {'content': [texts_conf["docstr"]["receive_act"]["text"]["agreement_head"], 
                texts_conf["docstr"]["receive_act"]["text"]["agreement_text"].format(name, cost) ], 'type': 'field'}], 
                texts_conf["docstr"]["receive_act"]["text"]["save_path"].format(date))
            return True
        elif reply == QtWidgets.QMessageBox.No:
            return None

    def check_func(self, product_name, name, phone):
        pattern_phone = r'[0-9]+'
        if product_name == '':
            QtWidgets.QMessageBox.warning(self, texts_conf["error"]["title"] , 
                texts_conf["error"]["text"]["ni_prod_name"])
            return None
        elif name == '':
            QtWidgets.QMessageBox.warning(self, texts_conf["error"]["title"], 
                texts_conf["error"]["text"]["ni_name"])
            return None
        elif phone == '':
            return None
        return True

    def search_func(self):                         #переход к окну отображения базы
        globals()['myapp'].hide()
        globals()['sew'].show()

    def keyPressEvent(self, e):
        if e.key() == QtCore.Qt.Key_Return:
            self.init_insert()


class SearchWin(MyWin):

    def __init__(self, parent=None):
        QtWidgets.QWidget.__init__(self, parent)
        self.ui = Ui_SearchWin()
        self.ui.setupUi(self)
        _translate = QtCore.QCoreApplication.translate
        items = []
        for i in enumerate(config):
            self.ui.comboBox_search_where.addItem("")
            self.ui.comboBox_search_where.setItemText(i[0], _translate("SearchWin", i[1]))
            items.append(i[1])
        self.ui.comboBox_search_where.setCurrentIndex(items.index('Принято'))
        self.ui.add_button.clicked.connect(self.add_item)
        self.asdc = 1
        self.order = 'Дата'
        self.where = 'Принято'
        self.create_default_table()

    def deletion(self):
        self.are_you_sure(self.deactivate, texts_conf["warning"]["text"]["delete"])

    def scaler(self):
        if self.where == self.ui.comboBox_search_where.currentText():
            self.create_default_table()
        else:
            self.define_where()
            self.ui.line_search.setText('')
            self.define_order_def()
            self.create_default_table()

    def define_where(self):
        self.where = self.ui.comboBox_search_where.currentText()
        
    def asc(self):
        self.asdc = 1
        self.from_last()

    def desc(self):
        self.asdc = -1
        self.from_last()

    def define_order_cont(self, order):
        self.order = order
        self.from_last()

    def define_order_def(self):
        self.order = config[self.where]['sorting_criteria'][0]
        self.create_default_table()

    def create_default_table(self):
        self.word = self.ui.line_search.text()
        datalist = self.get_data(self.word, self.where, self.order)
        self.create_table(datalist)

    def create_table(self, datalist):
        table_model = ApplicationsTableModel(datalist, config[self.where]['header'], self, table=self.where)
        self.ui.tableView_applications.setModel(table_model)
        self.ui.tableView_applications.resizeColumnsToContents()
        for i in config[self.where]['delegate_col']:
            self.ui.tableView_applications.setItemDelegateForColumn(i, Delegate(self,table=self.where))
        self.ui.tableView_applications.horizontalHeader().setStretchLastSection(True)

    def contextMenuEvent(self, event):

        data_opt = {None: lambda: 0+0}

        def assign_act(parent, string, func):
            return (parent.addAction(self.tr(string)), func)

        actionMenu = QtWidgets.QMenu(self.tr("&Actions"), self)
        (del_action, data_opt[del_action]) = assign_act(actionMenu, 
            'Удалить', lambda : self.are_you_sure(self.delete_item, texts_conf["warning"]["text"]["delete"]))

        sorting_criteria = QtWidgets.QMenu(self.tr("&Сортировать по"), self)
        sorting_order = QtWidgets.QMenu(self.tr("&Порядок"), self)

        (asc, data_opt[asc]) = assign_act(sorting_order, 'По возрастанию', self.asc)
        (desc, data_opt[desc]) = assign_act(sorting_order, 'По убыванию', self.desc)
        
        actionMenu.addMenu(sorting_order)

        if self.where in ['Принято', 'В работе', 'Завершено']:
            (set_status, data_opt[set_status]) = assign_act(actionMenu, 
                "Сменить статус", lambda: self.are_you_sure(self.move_to_next_status, 
                    'Вы точно хотите сменить \nстатус этой записи?'))

        if self.where in ['Принято', 'В работе', 'Завершено', 'Архив']:
            (show_all, data_opt[show_all]) = assign_act( actionMenu, 'Показать все сведения', self.show_all)

            point_status = QtWidgets.QMenu(self.tr("&Назначить статус"), self)

            (point_got, data_opt[point_got]) = assign_act(point_status,
                'Принято', lambda: self.set_status('Принято'))
            (point_in_work, data_opt[point_in_work]) = assign_act(point_status,
                'В работе', lambda: self.set_status('В работе'))
            (point_finished, data_opt[point_finished]) = assign_act(point_status,
                'Завершено', lambda: self.set_status('Завершено'))
            (point_archive, data_opt[point_archive]) = assign_act(point_status,
                'Архив', lambda: self.set_status('Архив'))

            date_search = QtWidgets.QMenu(self.tr("&Поиск по периоду"), self)

            (date_search_got, data_opt[date_search_got]) = assign_act(date_search,
                'Получено', lambda: self.date_search('Дата'))
            (date_search_done, data_opt[date_search_done]) = assign_act(date_search,
                'Завершено', lambda: self.date_search('Дата завершения'))
            (date_search_sent, data_opt[date_search_sent]) = assign_act(date_search,
                'Отправлено', lambda: self.date_search('Дата отгрузки'))

            actionMenu.addMenu(point_status)
            actionMenu.addMenu(date_search)

        if self.where in ['Завершено', 'Архив']:
            (final_docs, data_opt[final_docs]) = assign_act(actionMenu, 
                'Сформировать счёт и акт', self.final_docs)

        if self.where in 'Акты приёма в ремонт':
            (date_search_got, data_opt[date_search_got]) = assign_act(actionMenu, 'Поиск по дате', )
            (form_act, data_opt[form_act]) = assign_act(actionMenu, 'Сформировать документ', self.form_act)

        for i in config[self.where]['sorting_criteria']:
            b = i[:]
            locals()[b] = sorting_criteria.addAction(self.tr(i))
            data_opt[locals()[b]] = lambda i=i: self.define_order_cont(i)

        actionMenu.addMenu(sorting_criteria)

        action = actionMenu.exec_(self.mapToGlobal(event.pos()))
        data_opt[action]()

    def form_act(self):
        def get_and_form(id_):
            x= db.Checks.find_one({'_id': id_})
            a =[ x[i] for i in ['Дата', 'Принято от', 'Контакт', 'Принял', 
            'Примечания', 'Гарантийность', 'Изделия', 'Num']]
            self.form_check(*a)
        self.single_selection(get_and_form)

    def show_all(self):
        def show_all(id_):
            alist = ['ФИО','Телефон','Принял','Изделие','Примечания','Дата',
            'Производитель','Акт приёма','Гарантийность','Статус','Место',
            'Передал в ремонт','Передано работнику','Перечень работ','Дата завершения',
            'Цена','Дата отгрузки','Зарплата','Расчёт','Получатель', '_id']
            x = db.Applications.find_one({'_id': id_})
            strr = ''
            for i in alist:
                strr+= i + ': {}\n'.format(x[i])
            QtWidgets.QMessageBox.information(self, 'Сведения', strr)
        self.single_selection(show_all)


    def final_docs(self):

        date = str(QtCore.QDateTime.currentDateTime().date().toPyDate())

        def get_name(id_):
            x = config[self.where]['db'].find_one({'_id': id_})
            return x['ФИО']

        name = self.single_selection(get_name)

        def check_names(id_):
            x = config[self.where]['db'].find_one({'_id': id_})
            if name != x['ФИО']:
                QtWidgets.QMessageBox.warning(sew, texts_conf["error"]["title"], 
                    texts_conf["error"]["text"]["fin_doc_names"])
                return None
            else:
                return True

        for i in self.multiple_selection(check_names):
            if i != True:
                return None

        content_list_pays = [{'type':'field', 'content': ['Поставщик:', 'ИП Дик']}, {'type':'field',
            'content': ['Покупатель:', name]}]
        content_list_works = []
        l = []

        def receipt_data(id_):
            x = config[self.where]['db'].find_one({'_id': id_})
            clw = []
            l = []
            wks = sep_count(x['Перечень работ'])
            clw.append({'type': 'field', 'content': ['Список работ по изделию "{}"'.format(x['Изделие']), '']})
            clw.append({'type': 'table', 'content': [['№', 'Наименование', 'Стоимость, руб.', 'Количество', 'Сумма, руб.'],
                [[str(q[0]+1), str(q[1]), db.Jobs.find_one({'Наименование': q[1]})['Стоимость'], str(wks[q[1]][0]), 
                str(int(wks[q[1]][0]) * int(db.Jobs.find_one({'Наименование': q[1]})['Стоимость']))]for q in enumerate(wks)]]})

            clw.append({'type': 'field','content': ['Итого:', 
                str(sum([int(wks[w][0])*int(db.Jobs.find_one({'Наименование': w})['Стоимость']) for w in wks]))+' руб.']})
            l.append('Ремонт изделия "{}"'.format(x['Изделие']))
            l.append(sum([int(wks[w][0])*int(db.Jobs.find_one({'Наименование': w})['Стоимость']) for w in wks]))
            return (x['Изделие'], l, clw, id_)

        ths = []
        ids = []
        for i in self.multiple_selection(receipt_data):
            ths.append(i[0])
            l.append(i[1])
            ids.append(str(i[3]))
            for x in i[2]:
                content_list_works.append(x)
        str_ths = ', '.join(ths)
        str_ids = ','.join(ids)
        for i in enumerate(l):
            i[1].insert(0, i[0]+1)
        insa = [{'type': 'table', 'content': [['№', 'Товары/работы/услуги', 'Цена, руб.'], l]},{'type':'field', 'content':['Итого:', 
            str(sum([ i[2] for i in l]))+' руб.']},{'type':'field', 'content': '\n\n\n\n\n\n'}, {'type': 'field', 
            'content': ['ИП Дик', '_____________________________________{}'.format(str_date(date))]}]
        for i in insa:
            content_list_pays.append(i)

        cn = db.FinDocs.count_documents({'date': date[:4]})+1

        f = open('template.docx', 'rb')
        template = Document(f)
        f.close()

        pay = self.create_doc(texts_conf["docstr"]["bill"]["title"].format(cn, str_date(date)), content_list_pays, 
            texts_conf["docstr"]["bill"]["text"]["save_path"].format(date), template)
        work = self.create_doc(texts_conf["docstr"]["jobs_done"]["title"].format(str(cn), str_date(date)), 
            content_list_works,texts_conf["docstr"]["jobs_done"]["text"]["save_path"].format(date))

        db.FinDocs.insert_one({'date': date,'year':date[:4], 'Клиент': name, 'Изделия': str_ths, 'ids':str_ids})
        QtWidgets.QMessageBox.information(self, texts_conf["message"]["title"], texts_conf["message"]["text"]["final_docs"])


    def move_to_next_status(self):
        if self.where == 'Принято':
            self.set_status('В работе')
        elif self.where == 'В работе':
            self.set_status('Завершено')
        elif self.where == 'Завершено':
            self.set_status('Архив')

    def get_data(self, search_word='', search_where='Принято',
                 order='Дата'):
        find_func = config[search_where]['db'].find
        self.arg = {}
        self.arg_func = config[search_where]['default_argument']
        datalist = self.db_func_arg_decorator(self.arg,
                            search_word, 
                            find_func,
                            order,
                            config[search_where]['filling_func'],
                            self.arg_func,
                            config[search_where]['Status'])
        return datalist

    def db_func_arg_decorator(self, arg, search_word, db_func, order, filling_func, arg_func, status):
        datalist = []
        if search_word == '':
            arg['Статус'] = status
        else:
            arg = arg_func(search_word)
            arg['Статус'] = status
        for i in db_func(arg).sort(order, self.asdc):
            tempvar = filling_func(i)
            datalist.append(tempvar)
        return datalist
                    # написать индексирование для всех полей

    def add_item(self):
        self.define_where()
        if self.where in ['Принято', 'В работе', 'Завершено', 'Архив']:
            globals()['sew'].hide()
            globals()['myapp'].show()
        elif config[self.where]['fields_order']:
            win_title = "Форма ввода записей в базу {}".format(config[self.where]['db_name'])
            def func(args):
                args['Статус'] = 'active'
                insert_func = config[self.where]['db'].insert_one(args)
                self.scaler()
                MyWin.reset_completers(globals()['myapp'])
            EnterForm = create_form_class(func,
                                          config[self.where]['fields'],
                                          config[self.where]['fields_order'],
                                          win_title)
            self.enterform = EnterForm()
            self.enterform.show()
        else:
            QtWidgets.QMessageBox.information(self, texts_conf["error"]["title"], 
                texts_conf["error"]["text"]["db_enterform"])
            return None

    def date_search(self, field):
        self.define_where()
        self.word = self.ui.line_search.text()
        def func(args):
            search_f = config[self.where]['db'].find
            self.arg_func = lambda i: {'$text' : { '$search': i }, field:{"$gte": args['С'], "$lte": args['По']}}
            self.arg = {field:{"$gte": args['С'], "$lte": args['По']}}
            dl = self.db_func_arg_decorator(self.arg, self.word, 
                                    search_f,
                                    self.order,
                                    config[self.where]['filling_func'],
                                    self.arg_func,
                                    config[self.where]['Status'])
            self.create_table(dl)
        fls = { 'С': { 'disp_field': 'С',
                        'cu_income': 'date',
                        'editor': 'date'},
                'По':{ 'disp_field': 'По',
                        'cu_income': 'date',
                        'editor': 'date'}}
        fields_order = ['С', 'По']
        win_title = 'Поиск по полю "{}"'.format(field)
        EnterForm = create_form_class(func,
                                    fls,
                                    fields_order,
                                    win_title,)
        self.enterform = EnterForm()
        self.enterform.show()

    def from_last(self):
        search_f = config[self.where]['db'].find
        dl = self.db_func_arg_decorator(self.arg, self.word, 
                                    search_f,
                                    self.order,
                                    config[self.where]['filling_func'],
                                    self.arg_func,
                                    config[self.where]['Status'])
        self.create_table(dl)


    def are_you_sure(self, function, string):
        reply = QtWidgets.QMessageBox.question(self,'Внимание', string,QtWidgets.QMessageBox.Yes, QtWidgets.QMessageBox.No)
        if reply == QtWidgets.QMessageBox.Yes:
            function()
        elif reply == QtWidgets.QMessageBox.No:
            pass

    def set_status(self, string):
        def set_status(anid):
            config[self.where]['db'].find_one_and_update({'_id': anid}, {'$set':{'Статус': string}})
        self.single_selection(set_status)
        self.scaler()

    def single_selection(self, function):
        id_f = config[self.where]['id_f']
        index = self.ui.tableView_applications.currentIndex()
        id_ = ObjectId(index.sibling(index.row(), id_f).data())
        return function(id_)

    def deactivate(self):
        self.set_status('inactive')

    def delete_item(self):
        def delete_func(anid):
            config[self.where]['db'].delete_one({'_id': anid})
        self.multiple_selection(delete_func)
        self.scaler()

    def multiple_selection(self, function):
        id_f = config[self.where]['id_f']
        ids = sorted(set(ObjectId(index.sibling(index.row(), 
            id_f).data()) for index in self.ui.tableView_applications.selectedIndexes()))
        for anid in ids:
            r = function(anid)
            yield r

    def check_if_selected(self, ids):
        if ids:
            QtWidgets.QMessageBox.warning(self, texts_conf["error"]["title"], 
                texts_conf["error"]["text"]["selection_chk"])

    def keyPressEvent(self, e):
        if e.key() == QtCore.Qt.Key_Return:
            self.scaler()
        elif e.key() == QtCore.Qt.Key_Delete:
            self.are_you_sure(self.deactivate, texts_conf["warning"]["text"]["delete"])


class ApplicationsTableModel(QtCore.QAbstractTableModel):
    def __init__(self, datain, headerdata, parent=None, table='Принято'):
        # Datain is list of lists tabledata, headerdata is a string list
        QtCore.QAbstractTableModel.__init__(self, parent)
        self.arraydata = datain
        self.headerdata = headerdata
        self.table = table

    def rowCount(self, parent):
        return len(self.arraydata)

    def columnCount(self, parent):
        return len(self.headerdata)

    def headerData(self, section, orientation, role):
        if orientation == QtCore.Qt.Horizontal and role == QtCore.Qt.DisplayRole:
            return QtCore.QVariant(self.headerdata[section])
        elif orientation == QtCore.Qt.Vertical and role == QtCore.Qt.DisplayRole:
            return QtCore.QAbstractTableModel.headerData(self, section,orientation, role)

    def data(self, index, role):
        if not index.isValid():
            return QtCore.QVariant()
        elif role != QtCore.Qt.DisplayRole:
            return QtCore.QVariant()
        return QtCore.QVariant(self.arraydata[index.row()][index.column()])

    def setData(self, index, value, role):
        if index.isValid() and role == QtCore.Qt.EditRole:
            self.arraydata[index.row()][index.column()] = QtCore.QVariant(value)
            self.emit(QtCore.SIGNAL("dataChanged(QModelIndex, QModelIndex)"),index, index)
            return True
        else:
            return False

    def flags(self, index):
        if index.column() in config[self.table]['delegate_col']:  #List of editable columns
            return QtCore.Qt.ItemIsEditable | QtCore.Qt.ItemIsEnabled | \
                   QtCore.Qt.ItemIsSelectable
        else:
            return QtCore.Qt.ItemIsEnabled | QtCore.Qt.ItemIsSelectable


class Delegate(QtWidgets.QStyledItemDelegate):
    def __init__(self, parent, table='Принято'):
        QtWidgets.QStyledItemDelegate.__init__(self, parent)
        self.table = table

    def createEditor(self, parent, options, index):
        cur_field = config[self.table]['editing_fs'][index.column()]
        source_list = None
        if cur_field['editor'] in ['c_line', 'combo']:
            source_list = cur_field['c_source']
        return conf_crf[cur_field['editor']](parent, source_list)

    def setEditorData(self, editor, index):
        a = config[self.table]['editing_fs'][index.column()]['editor']
        if a=='c_line':
            a = a[2:]
        conf_set[a](editor, index)

    def setModelData(self, editor, model, index):
        cur_field = config[self.table]['editing_fs'][index.column()]
        a = config[self.table]['editing_fs'][index.column()]['editor']
        if a == 'c_line':
            a = a[2:]
        new_data = conf_stm[a](editor)
        ident = model.index(index.row(),config[self.table]['id_f']).data()
        if 'cu_income' in cur_field:
            new_data = cu_conf[cur_field['cu_income']](new_data, ident)
        config[self.table]['db'].find_one_and_update({'_id': ObjectId(ident)},
            { '$set': {cur_field['field']: new_data}})
        sew.from_last()


class Completer(QtWidgets.QCompleter):
    def __init__(self, parent=None):
        super(Completer, self).__init__(parent)
        self.setCaseSensitivity(QtCore.Qt.CaseInsensitive)
        self.setCompletionMode(QtWidgets.QCompleter.PopupCompletion)
        self.setWrapAround(False)

    # Add texts instead of replace
    def pathFromIndex(self, index):
        path = QtWidgets.QCompleter.pathFromIndex(self, index)
        lst = str(self.widget().text()).split(',')
        if len(lst) > 1:
            path = '%s, %s' % (','.join(lst[:-1]), path)
        return path

    # Add operator to separate between texts
    def splitPath(self, path):
        path = str(path.split(',')[-1]).lstrip(' ')
        return [path]

def create_form_class(func, field_dict, fields_order, win_title):
    class EnterForm(MyWin):
        def __init__(self):
            QtWidgets.QWidget.__init__(self)
            self.ui = Ui_Form()
            self.ui.setupUi(self)
            self.ui.win_title = win_title

            counter = 0
            for i in fields_order:
                self.ui.create_fields(self, counter, conf_crf[field_dict[i]['editor']](self))
                counter+=1

            self.ui.creating_buttons(self, counter)

            disp_fields = []
            for i in fields_order:
                disp_fields.append(field_dict[i]['disp_field'])

            self.ui.start_retranslate(self, disp_fields)
            self.ui.buttonBox.clicked.connect(self.handleButtonClick)

        def handleButtonClick(self, button):
            sb = self.ui.buttonBox.standardButton(button)
            if sb == QtWidgets.QDialogButtonBox.Ok:
                self.check_func()
            elif sb == QtWidgets.QDialogButtonBox.Cancel:
                self.close()

        def check_func(self):
            counter = 0
            args = {}
            for i in fields_order:
                a = 'field_' + str(counter)
                data = conf_stm[field_dict[i]['editor']](self.ui.__dict__[a])
                if self.empty_field_check(data, i):
                    if 'cu_income' in field_dict[i]:
                        if not cu_conf[field_dict[i]['cu_income']](data, None):
                            return None
                        args[i] = cu_conf[field_dict[i]['cu_income']](data, None)
                    else:
                        args[i] = data
                else:
                    return None
                counter+=1
            func(args)
            return self.close()

        def empty_field_check(self, data, field):
            if data == '':
                msg = 'Не заполнено поле "{}"'.format(field)
                QtWidgets.QMessageBox.warning(self, 'Ошибка', msg)
                return None
            else:
                return data

        def keyPressEvent(self, e):
            if e.key() == QtCore.Qt.Key_Escape:
                self.close()
            elif e.key() == QtCore.Qt.Key_Return:
                self.check_func()

    return EnterForm

def sep_count(string):
    pattern = r'[(](?=\d+[)])'
    s = string.split(',')
    b =[]
    for i in s:
        b.append(i.strip())
    adic = {}
    for i in b:
        if re.match(r'.+[(]\d+[)]', i):
            adic[re.split(pattern, i)[0].strip()] = [abs(int(re.split(pattern, i)[1].strip('()')))]
        else:
            adic[i.strip()] = [1]
    # adic = {re.split(pattern, i)[0].strip(): re.split(pattern, i)[1].strip('()') for i in b}
    return adic


if __name__ == "__main__":
    client()
    bdir = os.getcwd()
    with open(os.path.join(bdir, 'texts_and_fields.json')) as texts_and_fields:
        texts_conf = json.load(texts_and_fields)
    app = QtWidgets.QApplication(sys.argv)
    myapp = MyWin()
    sew = SearchWin()
    myapp.show()
    sys.exit(app.exec_())